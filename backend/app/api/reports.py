"""报告生成 API 路由"""
from __future__ import annotations
import re
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from app.database import get_db
from app.models.report import Report
from app.schemas.report import (
    ReportGenerateInput, ReportSummary, ReportDetail, ReportHistoryResponse,
    VALID_ACTIVITY_TYPES,
)
from app.schemas.intent import IntentInput
from app.agents.agent_00_intent_recognizer.intent_recognizer import recognize_intent
from app.agents.agent_01_orchestrator.orchestrator import orchestrate
from app.schemas.orchestrator import OrchestratorInput
from app.services.feishu_data import get_feishu_products
from app.services.data_generator import get_demo_products
from app.utils.auth import require_auth
import urllib.parse

router = APIRouter(dependencies=[Depends(require_auth)], prefix="/api/v1/reports", tags=["报告"])

ACTIVITY_LABELS = {
    "double11": "双11大促",
    "618": "618大促",
    "new_product": "新品发布",
    "clearance": "清仓促销",
    "daily": "日常促销",
    "flash_sale": "限时秒杀",
    "member_day": "会员日",
    "festival": "节日促销",
    "pre_sale": "预售活动",
    "group_buy": "拼团活动",
    "anniversary": "店庆活动",
    "season_change": "换季清仓",
}


@router.post("/generate")
def generate_report(input_data: ReportGenerateInput, db: Session = Depends(get_db)):
    """生成选品分析报告：用户选择商品 + 活动类型 -> 多智能体协作 -> 产出报告"""
    # 校验活动类型
    if input_data.activity_type not in VALID_ACTIVITY_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"无效的活动类型: {input_data.activity_type}，可选: {', '.join(sorted(VALID_ACTIVITY_TYPES))}"
        )

    all_products = get_feishu_products()
    if not all_products:
        all_products = get_demo_products("", count=40)

    selected_products = [p for p in all_products if str(p["product_id"]) in input_data.product_ids]
    if not selected_products:
        raise HTTPException(status_code=400, detail="未找到匹配的商品，请检查商品ID")

    from collections import Counter
    categories = [p.get("category", "") for p in selected_products if p.get("category")]
    main_category = Counter(categories).most_common(1)[0][0] if categories else "综合"

    product_summary = "\n".join(
        f"- {p['title']}（ID: {p['product_id']}, 售价: {p.get('current_price', p.get('price', 0))}, 库存: {p.get('current_stock', '?')}）"
        for p in selected_products
    )
    activity_label = ACTIVITY_LABELS[input_data.activity_type]

    user_message = (
        f"请对以下 {len(selected_products)} 个商品进行综合分析，活动类型为「{activity_label}」。\n\n"
        f"商品列表：\n{product_summary}\n\n"
        f"要求：\n"
        f"1. 选品分析 —— 评估各商品爆款潜力，排序推荐\n"
        f"2. 定价策略 —— 结合活动类型给出最优定价建议\n"
        f"3. 竞品分析 —— 与市场上同类商品对比\n"
        f"4. 营销文案 —— 为本次活动生成推广文案\n"
        f"5. 活动策划 —— 制定完整活动方案\n"
        f"6. 补货建议 —— 根据库存和活动预期给出补货/清仓建议"
    )

    intent_input = IntentInput(
        user_message=user_message,
        session_id=input_data.session_id,
        turn_number=1,
    )
    try:
        intent_result = recognize_intent(intent_input)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"意图识别失败: {str(e)}")

    # 安全获取 for_downstream
    intent_downstream = (
        intent_result.for_downstream.copy()
        if intent_result.for_downstream
        else {}
    )
    intent_downstream["required_tasks"] = [
        "product_selection", "pricing_strategy", "competitor_analysis",
        "marketing_copy", "promotion_plan", "inventory_advice",
    ]
    intent_downstream["category"] = main_category
    intent_downstream["top_n"] = len(selected_products)

    orch_input = OrchestratorInput(
        intent_result={
            "parsed_result": intent_result.parsed_result.model_dump(),
            "for_downstream": intent_downstream,
            "for_display": intent_result.for_display or "",
        },
        session_id=input_data.session_id,
    )
    try:
        orch_result = orchestrate(orch_input)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"智能体编排失败: {str(e)}")

    report_content = {
        "activity_type": input_data.activity_type,
        "activity_label": activity_label,
        "category": main_category,
        "products": [
            {
                "product_id": p["product_id"],
                "title": p["title"],
                "price": p.get("current_price", p.get("price", 0)),
                "category": p.get("category", ""),
                "stock": p.get("current_stock", 0),
            }
            for p in selected_products
        ],
        "orchestrator_result": orch_result.model_dump(),
        "task_plan": orch_result.task_plan.model_dump() if orch_result.task_plan else {},
        "phase_results": [pr.model_dump() for pr in orch_result.phase_results] if orch_result.phase_results else [],
        "final_report": orch_result.final_report,
    }

    final_report = orch_result.final_report or {}
    summary = final_report.get("summary", "") if isinstance(final_report, dict) else ""

    try:
        report = Report(
            session_id=input_data.session_id,
            product_ids=input_data.product_ids,
            activity_type=input_data.activity_type,
            category=main_category,
            summary=summary,
            report_content=report_content,
        )
        db.add(report)
        db.commit()
        db.refresh(report)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"保存报告失败: {str(e)}")

    return {
        "report_id": report.id,
        "session_id": report.session_id,
        "activity_type": report.activity_type,
        "activity_label": activity_label,
        "category": report.category,
        "product_count": len(selected_products),
        "summary": report.summary,
        "report_content": report_content,
        "created_at": report.created_at.isoformat(),
    }


@router.get("/history", response_model=ReportHistoryResponse)
def list_reports(skip: int = 0, limit: int = 20, db: Session = Depends(get_db)):
    """获取历史报告列表（按时间倒序）"""
    total = db.query(Report).count()
    reports = (
        db.query(Report)
        .order_by(Report.created_at.desc())
        .offset(skip)
        .limit(limit)
        .all()
    )
    items = [
        ReportSummary(
            id=r.id,
            session_id=r.session_id,
            product_count=len(r.product_ids) if r.product_ids else 0,
            activity_type=r.activity_type,
            category=r.category,
            summary=r.summary,
            created_at=r.created_at.isoformat(),
        )
        for r in reports
    ]
    return ReportHistoryResponse(reports=items, total=total)


@router.get("/{report_id}", response_model=ReportDetail)
def get_report(report_id: str, db: Session = Depends(get_db)):
    """获取单份报告详情"""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")
    return ReportDetail(
        id=report.id,
        session_id=report.session_id,
        product_ids=report.product_ids or [],
        activity_type=report.activity_type,
        category=report.category,
        summary=report.summary,
        report_content=report.report_content or {},
        created_at=report.created_at.isoformat(),
    )


# === 追问对话 ===
from pydantic import BaseModel, Field


class ReportChatInput(BaseModel):
    user_message: str = Field(..., min_length=1, description="用户追问内容")
    conversation_history: list[dict] = Field(default_factory=list, description="之前的对话历史")


@router.post("/{report_id}/chat")
def chat_followup(report_id: str, input_data: ReportChatInput, db: Session = Depends(get_db)):
    """基于报告上下文进行追问对话"""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    report_content = report.report_content or {}
    products = report_content.get("products", [])
    activity_label = report_content.get("activity_label", report.activity_type)
    category = report.category or "综合"
    summary = report.summary or ""

    if products:
        product_context = "\n".join(
            f"- {p.get('title', '未知')}（{p.get('price', '?')}，{p.get('category', '?')}）"
            for p in products[:10]
        )
    else:
        product_context = "（无具体商品数据）"

    context = (
        f"你是一个电商选品分析助手。用户刚才生成了以下分析报告：\n\n"
        f"活动类型：{activity_label}\n"
        f"类目：{category}\n"
        f"涉及商品：\n{product_context}\n\n"
        f"报告摘要：{summary[:1000]}\n\n"
        f"现在用户追问：{input_data.user_message}\n\n"
        f"请基于以上报告上下文，用中文简洁回答用户的问题。直接输出最终回答，不要输出思考过程。"
        f"如果是具体的数据/建议，尽量引用报告中的内容。"
    )

    from app.utils.llm_client import chat_completion
    try:
        reply = chat_completion(
            "你是一个专业的电商选品分析助手，回答基于已有分析报告，语气专业友好。直接回答，不要输出思考过程。",
            context,
            temperature=0.4,
            max_tokens=1024,
        )
        reply = reply.strip()
        reply = re.sub(r"<think>.*?</think>", "", reply, flags=re.DOTALL).strip()
        if not reply:
            reply = "抱歉，未能生成有效回复，请重新提问。"
    except Exception as e:
        reply = f"抱歉，处理追问时出错：{str(e)}"

    history = list(input_data.conversation_history or [])
    history.append({"role": "user", "content": input_data.user_message})
    history.append({"role": "assistant", "content": reply})

    if len(history) > 20:
        history = history[-20:]

    return {"reply": reply, "conversation_history": history}

# === 报告导出 ===
from io import BytesIO
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


@router.get("/{report_id}/export")
def export_report_docx(report_id: str, db: Session = Depends(get_db)):
    """导出报告为 .docx 文档"""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

# === Markdown 导出 ===
AGENT_LABELS_MD = {
    "product_selection": "选品分析",
    "trend_forecast": "趋势预测",
    "competitor_analysis": "竞品分析",
    "user_profile": "用户画像",
    "pricing_strategy": "定价策略",
    "marketing_copy": "营销文案",
    "inventory_advice": "补货/清仓建议",
    "promotion_plan": "活动策划",
}


def _report_to_markdown(report) -> str:
    """将报告对象转换为 Markdown 字符串"""
    lines = []
    rc = report.report_content or {}
    activity_label = rc.get("activity_label", report.activity_type)
    category = report.category or "综合"
    products = rc.get("products", [])

    lines.append("# 电商选品分析报告")
    lines.append("")
    lines.append(f"**分析类目**: {category}  ")
    lines.append(f"**活动类型**: {activity_label}  ")
    lines.append(f"**分析商品数**: {len(report.product_ids or [])} 个  ")
    dt = report.created_at.strftime("%Y-%m-%d %H:%M") if report.created_at else ""
    lines.append(f"**生成时间**: {dt}")
    lines.append("")

    lines.append("## 商品清单")
    lines.append("")
    if products:
        lines.append("| 商品ID | 商品名称 | 类目 | 售价 | 库存 |")
        lines.append("|--------|----------|------|------|------|")
        for p in products:
            pid = p.get("product_id", "")
            title = p.get("title", "")
            cat = p.get("category", "")
            price = p.get("price", "")
            stock = p.get("stock", "")
            lines.append(f"| {pid} | {title} | {cat} | {price} | {stock} |")
    else:
        lines.append("（无商品数据）")
    lines.append("")

    if report.summary:
        lines.append("## 报告摘要")
        lines.append("")
        lines.append(report.summary)
        lines.append("")

    lines.append("## 智能体分析详情")
    lines.append("")

    orch = rc.get("orchestrator_result", {})
    phases = orch.get("phase_results", []) or rc.get("phase_results", [])

    if phases:
        for phase in phases:
            phase_num = phase.get("phase", "?")
            lines.append(f"### 第 {phase_num} 阶段")
            lines.append("")
            for agent in (phase.get("agents", []) or []):
                task_type = agent.get("task_type", agent.get("agent_name", "未知"))
                label = AGENT_LABELS_MD.get(task_type, task_type)
                lines.append(f"#### {label}")
                lines.append("")

                summary = agent.get("summary", "")
                if summary:
                    lines.append(summary)
                else:
                    output = agent.get("output", agent.get("result", ""))
                    payload = {}
                    if isinstance(output, dict):
                        payload = output.get("payload", {})
                    if payload:
                        for_display = payload.get("for_display", "")
                        if for_display:
                            lines.append(str(for_display))
                        else:
                            for k, v in payload.items():
                                if k in ("for_downstream", "agent_name", "envelope"):
                                    continue
                                if v is not None and v != "":
                                    sv = str(v)
                                    lines.append(f"- **{k}**: {sv[:500]}")
                    elif isinstance(output, dict):
                        for k, v in output.items():
                            if k == "envelope":
                                continue
                            if v is not None and v != "":
                                lines.append(f"- **{k}**: {str(v)[:300]}")
                    elif isinstance(output, str) and output:
                        lines.append(output)
                    elif isinstance(output, list):
                        for item in output:
                            if isinstance(item, str):
                                lines.append(f"- {item}")

                recommendation = agent.get("recommendation", "")
                if recommendation:
                    lines.append(f"> {recommendation}")
                lines.append("")
    else:
        final_report = orch.get("final_report", "")
        for_display = orch.get("for_display", "")
        if for_display:
            lines.append(for_display)
        elif final_report:
            lines.append(str(final_report))
        else:
            lines.append("（暂无明显数据）")

    return "\n".join(lines)


@router.get("/{report_id}/export/md")
def export_report_markdown(report_id: str, db: Session = Depends(get_db)):
    """导出报告为 Markdown 文件"""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    md_content = _report_to_markdown(report)

    category = report.category or "综合"
    activity_label = (report.report_content or {}).get("activity_label", report.activity_type)
    dt = report.created_at.strftime("%Y%m%d") if report.created_at else ""
    filename = f"分析报告_{category}_{activity_label}_{dt}.md"

    return StreamingResponse(
        BytesIO(md_content.encode("utf-8")),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}"},
    )



    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading("电商选品分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    activity_label = (report.report_content or {}).get("activity_label", report.activity_type)
    category = report.category or "综合"

    info_table = doc.add_table(rows=4, cols=2, style="Table Grid")
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("分析类目", category),
        ("活动类型", activity_label),
        ("分析商品数", f"{len(report.product_ids or [])} 个"),
        ("生成时间", report.created_at.strftime("%Y-%m-%d %H:%M") if report.created_at else ""),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.cell(i, 0).text = k
        info_table.cell(i, 1).text = v
        for cell in [info_table.cell(i, 0), info_table.cell(i, 1)]:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # 商品清单
    doc.add_heading("商品清单", level=1)
    products = (report.report_content or {}).get("products", [])
    if products:
        pt = doc.add_table(rows=len(products) + 1, cols=5, style="Table Grid")
        pt.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["商品ID", "商品名称", "类目", "售价", "库存"]
        for j, h in enumerate(headers):
            pt.cell(0, j).text = h
            for p in pt.cell(0, j).paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for i, prod in enumerate(products):
            pt.cell(i + 1, 0).text = str(prod.get("product_id", ""))
            pt.cell(i + 1, 1).text = str(prod.get("title", ""))
            pt.cell(i + 1, 2).text = str(prod.get("category", ""))
            pt.cell(i + 1, 3).text = str(prod.get("price", ""))
            pt.cell(i + 1, 4).text = str(prod.get("stock", ""))

    doc.add_paragraph()

    # 报告摘要
    doc.add_heading("报告摘要", level=1)
    if report.summary:
        doc.add_paragraph(report.summary)
    else:
        doc.add_paragraph("（无摘要）")

    # Agent 分析结果
    doc.add_heading("智能体分析详情", level=1)
    orchestrator = (report.report_content or {}).get("orchestrator_result", {})
    phase_results = orchestrator.get("phase_results", [])
    if not phase_results:
        phase_results = (report.report_content or {}).get("phase_results", [])

    if phase_results:
        for phase in phase_results:
            phase_num = phase.get("phase", "?")
            doc.add_heading(f"第 {phase_num} 阶段", level=2)
            agents_in_phase = phase.get("agents", [])
            for agent in agents_in_phase:
                agent_name = agent.get("task_type", agent.get("agent_name", "未知"))
                label_map = {
                    "product_selection": "选品分析",
                    "trend_forecast": "趋势预测",
                    "competitor_analysis": "竞品分析",
                    "user_profile": "用户画像",
                    "pricing_strategy": "定价策略",
                    "marketing_copy": "营销文案",
                    "inventory_advice": "补货/清仓建议",
                    "promotion_plan": "活动策划",
                }
                display = label_map.get(agent_name, agent_name)
                doc.add_heading(display, level=3)

                output = agent.get("output", agent.get("result", ""))
                if isinstance(output, dict):
                    for k, v in output.items():
                        if isinstance(v, (str, int, float)) and v:
                            doc.add_paragraph(f"{k}: {v}")
                elif isinstance(output, str) and output:
                    doc.add_paragraph(output)
                elif isinstance(output, list):
                    for item in output:
                        if isinstance(item, str):
                            doc.add_paragraph(item, style="List Bullet")

                recommendation = agent.get("recommendation", "")
                if recommendation:
                    doc.add_paragraph(f"建议: {recommendation}")
    else:
        final_report = orchestrator.get("final_report", {})
        for_display = orchestrator.get("for_display", "")
        if for_display:
            doc.add_paragraph(for_display)
        elif final_report:
            doc.add_paragraph(str(final_report))

    # 保存到内存
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"分析报告_{category}_{activity_label}_{report.created_at.strftime('%Y%m%d') if report.created_at else ''}.docx"
    encoded_filename = filename.encode("utf-8").decode("latin-1", errors="ignore") or "report.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}"},
    )

# === PDF 导出 ===
from fpdf import FPDF


class ChinesePDF(FPDF):
    def __init__(self):
        super().__init__()
        # Add a built-in Unicode font
        self.add_font("CJK", "", r"C:\Windows\Fonts\simsun.ttc", uni=True)
        self.add_font("CJK", "B", r"C:\Windows\Fonts\simsunb.ttf", uni=True)


@router.get("/{report_id}/export/pdf")
def export_report_pdf(report_id: str, db: Session = Depends(get_db)):
    """导出报告为 PDF 文件"""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    rc = report.report_content or {}
    activity_label = rc.get("activity_label", report.activity_type)
    category = report.category or "综合"
    products = rc.get("products", [])
    dt = report.created_at.strftime("%Y-%m-%d %H:%M") if report.created_at else ""

    pdf = ChinesePDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("CJK", "B", 20)
    pdf.cell(0, 14, "电商选品分析报告", ln=True, align="C")
    pdf.ln(6)

    # Info
    pdf.set_font("CJK", "", 11)
    pdf.cell(0, 8, f"分析类目: {category}    活动类型: {activity_label}    商品数: {len(report.product_ids or [])} 个    时间: {dt}", ln=True)
    pdf.ln(6)

    # Products table
    if products:
        pdf.set_font("CJK", "B", 13)
        pdf.cell(0, 10, "商品清单", ln=True)
        pdf.ln(2)
        pdf.set_font("CJK", "B", 10)
        col_w = [18, 65, 30, 30, 30]
        headers = ["ID", "商品名称", "类目", "售价", "库存"]
        for j, (h, w) in enumerate(zip(headers, col_w)):
            pdf.cell(w, 8, h, border=1, align="C")
        pdf.ln()
        pdf.set_font("CJK", "", 10)
        for p in products:
            vals = [str(p.get("product_id","")), str(p.get("title",""))[:18], str(p.get("category",""))[:8], str(p.get("price","")), str(p.get("stock",""))]
            for v, w in zip(vals, col_w):
                pdf.cell(w, 8, v, border=1)
            pdf.ln()
        pdf.ln(6)

    # Summary
    if report.summary:
        pdf.set_font("CJK", "B", 13)
        pdf.cell(0, 10, "报告摘要", ln=True)
        pdf.ln(2)
        pdf.set_font("CJK", "", 10)
        pdf.multi_cell(0, 6, report.summary[:2000])
        pdf.ln(4)

    # Agent analysis
    orch = rc.get("orchestrator_result", {})
    phases = orch.get("phase_results", []) or rc.get("phase_results", [])
    if phases:
        pdf.set_font("CJK", "B", 13)
        pdf.cell(0, 10, "智能体分析详情", ln=True)
        pdf.ln(2)
        for phase in phases:
            pdf.set_font("CJK", "B", 11)
            pdf.cell(0, 8, f"第 {phase.get('phase','?')} 阶段", ln=True)
            for agent in (phase.get("agents", []) or []):
                task_type = agent.get("task_type", agent.get("agent_name", ""))
                label = AGENT_LABELS_MD.get(task_type, task_type)
                pdf.set_font("CJK", "B", 10)
                pdf.cell(0, 7, f"  {label}", ln=True)
                pdf.set_font("CJK", "", 10)
                summary = agent.get("summary", "")
                if summary:
                    pdf.multi_cell(0, 6, summary[:800])
                pdf.ln(2)
    else:
        fd = orch.get("for_display", orch.get("final_report", ""))
        if fd:
            pdf.set_font("CJK", "", 10)
            pdf.multi_cell(0, 6, str(fd)[:3000])

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)

    filename = f"分析报告_{category}_{activity_label}_{report.created_at.strftime('%Y%m%d') if report.created_at else ''}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}"},
    )

